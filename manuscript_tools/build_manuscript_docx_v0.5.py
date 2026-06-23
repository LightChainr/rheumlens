from pathlib import Path
import re
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "RheumLens_manuscript_working_v0.5.md"
OUTPUT = ROOT / "RheumLens_manuscript_working_v0.5.docx"

TITLE = (
    "RheumLens: a donor-level audit benchmark for single-cell foundation-model "
    "disease representations in systemic lupus erythematosus"
)

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)
LIGHT_FILL = "F4F6F9"


def set_font(run, name="Calibri", size=11, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_inline(paragraph, text):
    text = text.replace(r"\[", "[").replace(r"\]", "]").replace(r"\~", "~")
    token_re = re.compile(r"(\*\*.*?\*\*|`.*?`|\*[^*]+?\*)")
    cursor = 0
    for match in token_re.finditer(text):
        if match.start() > cursor:
            set_font(paragraph.add_run(text[cursor:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), bold=True)
        elif token.startswith("`"):
            set_font(paragraph.add_run(token[1:-1]), name="Courier New", size=9.5)
        else:
            set_font(paragraph.add_run(token[1:-1]), italic=True)
        cursor = match.end()
    if cursor < len(text):
        set_font(paragraph.add_run(text[cursor:]))


def add_body(doc, text, style=None, keep_with_next=False, size=11, after=8, line_spacing=1.333):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.keep_with_next = keep_with_next
    add_inline(p, text)
    if size != 11:
        for run in p.runs:
            run.font.size = Pt(size)
    return p


def add_html_table(doc, block, caption=None):
    root = ET.fromstring(block)
    rows = []
    for tr in root.findall(".//tr"):
        vals = ["".join(td.itertext()).strip() for td in list(tr)]
        if vals and not all(re.fullmatch(r"[-:]+", v or "") for v in vals):
            rows.append(vals)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        add_inline(p, caption)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    widths = [9360 // ncols] * ncols
    if ncols == 5:
        widths = [1650, 1550, 1500, 1400, 3260]
    elif ncols == 4:
        widths = [2450, 2100, 2350, 2460]
    set_table_geometry(table, widths)
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            value = row[j] if j < len(row) else ""
            run = p.add_run(value)
            set_font(run, size=9.5, bold=(i == 0))
            if i == 0:
                set_cell_shading(cell, LIGHT_FILL)
            if j > 0 and len(value) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_markdown_table(doc, rows, caption=None):
    parsed = []
    for line in rows:
        line = line.strip()
        if not line.startswith("|"):
            continue
        vals = [v.strip() for v in line.strip("|").split("|")]
        if vals and not all(re.fullmatch(r"[-: ]+", v or "") for v in vals):
            parsed.append(vals)
    if not parsed:
        return
    ncols = max(len(r) for r in parsed)
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        add_inline(p, caption)
    table = doc.add_table(rows=len(parsed), cols=ncols)
    table.style = "Table Grid"
    widths = [max(900, 9360 // ncols)] * ncols
    set_table_geometry(table, widths)
    for i, row in enumerate(parsed):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            value = row[j] if j < len(row) else ""
            run = p.add_run(value)
            set_font(run, size=8.4, bold=(i == 0))
            if i == 0:
                set_cell_shading(cell, LIGHT_FILL)
            if j > 0 and len(value) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_markdown_image(doc, line):
    match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
    if not match:
        return False
    alt, path = match.group(1), match.group(2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    try:
        run = p.add_run()
        run.add_picture(path, width=Inches(6.45))
    except Exception as exc:
        add_inline(p, f"[Figure image could not be embedded: {path}; {exc}]")
    if alt:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        set_font(cap.add_run(alt), size=9, italic=True, color=MUTED)
    return True


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("RheumLens | Working manuscript v0.5"), size=9, color=MUTED)
    page_number(sec.footer.paragraphs[0])


def build():
    raw = SOURCE.read_text(encoding="utf-8")
    raw = raw.split("## Supplementary material plan", 1)[0].strip()
    raw = re.sub(r'\n<page url=.*?</page>\s*$', '', raw, flags=re.S)

    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True
    set_font(title.add_run(TITLE), size=19, bold=True, color=DARK_BLUE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(16)
    set_font(meta.add_run("Working draft v0.5 · 23 June 2026 · Target journal to be selected"), size=10, italic=True, color=MUTED)

    warning = doc.add_table(rows=1, cols=1)
    warning.style = "Table Grid"
    set_table_geometry(warning, [9360])
    set_cell_shading(warning.cell(0, 0), "FFF4CE")
    p = warning.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run("EDITORIAL STATUS — NOT FOR SUBMISSION: "), size=9.5, bold=True)
    set_font(p.add_run("Formal P5/P6/P8/P9 integration, figures and supplementary table package are assembled. Author metadata, journal formatting, repository links and reference audit remain pending."), size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    pending_caption = None
    lines = raw.splitlines()
    i = 0
    current_section = ""
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("<table>"):
            block = [line]
            while "</table>" not in block[-1] and i + 1 < len(lines):
                i += 1
                block.append(lines[i])
            add_html_table(doc, "\n".join(block), pending_caption)
            pending_caption = None
        elif line.startswith("|"):
            block = [line]
            while i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
                i += 1
                block.append(lines[i].strip())
            add_markdown_table(doc, block, pending_caption)
            pending_caption = None
        elif line.startswith("!["):
            add_markdown_image(doc, line)
        elif line.startswith("## "):
            current_section = line[3:]
            p = doc.add_paragraph(line[3:], style="Heading 1")
            p.paragraph_format.keep_with_next = True
        elif line.startswith("### "):
            p = doc.add_paragraph(line[4:], style="Heading 2")
            p.paragraph_format.keep_with_next = True
        elif re.match(r"\*\*Table \d+\.", line):
            pending_caption = line
        elif re.match(r"^[-*] ", line):
            add_body(doc, line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            add_body(doc, re.sub(r"^\d+\. ", "", line), style="List Number")
        else:
            if current_section == "References":
                add_body(doc, line, size=9.5, after=3, line_spacing=1.10)
            elif current_section == "Figure legends":
                add_body(doc, line, size=10, after=5, line_spacing=1.15)
            else:
                add_body(doc, line)
        i += 1

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Working manuscript for Bioinformatics Advances"
    doc.core_properties.keywords = "RheumLens; scGPT; SLE; donor-level classification"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
